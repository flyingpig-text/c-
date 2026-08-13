# -*- coding: utf-8 -*-
"""Convert a simple Chinese markdown document (headings, tables, paragraphs) to docx."""

from pathlib import Path
import sys

from docx import Document


def add_md_table(doc, rows):
    clean = lambda s: s.replace("`", "").replace("**", "").strip()
    header = [clean(c) for c in rows[0].strip().strip("|").split("|")]
    data = [
        [clean(c) for c in row.strip().strip("|").split("|")]
        for row in rows[1:]
        if not all(set(ch) <= set("-: ") for ch in row.strip().strip("|").split("|"))
    ]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, cell_text in enumerate(header):
        table.rows[0].cells[i].text = cell_text
    for row in data:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            if i < len(cells):
                cells[i].text = cell_text


def main():
    src = Path(__file__).with_name("水下服务器热设计参数与算法交付清单_数据补全.md")
    default = Path(__file__).with_name("水下服务器热设计参数与算法交付清单_数据补全.docx")
    dst = Path(sys.argv[1]) if len(sys.argv) > 1 else default
    doc = Document()
    lines = src.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i])
                i += 1
            add_md_table(doc, rows)
            continue
        elif line.strip() == "" or line.strip() == "---":
            pass
        else:
            doc.add_paragraph(line.replace("`", "").replace("**", ""))
        i += 1
    doc.save(dst)
    print(dst)


if __name__ == "__main__":
    main()
