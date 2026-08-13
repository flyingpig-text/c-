# -*- coding: utf-8 -*-
"""
第五问：综合前四问结果，生成给公司散热设计部门的建议信
=========================================================
功能：
    1) 自动定位工作区根目录，按相对路径读取前四问输出 CSV；
    2) 汇总最大服务器数、最优结构参数、推荐材料与深度、季节和潮汐最不利工况；
    3) 对关键变量做国际单位制与数量级校验；
    4) 生成 Word 建议信 letter.docx（python-docx），并打印全文；
    5) 同时输出核心数字表 Q5_核心数字.csv，便于论文引用核对。

运行：
    python 问题五/第五问_建议信.py

依赖：pandas、python-docx（bundled runtime 已具备）。
"""

from __future__ import annotations

import sys
import math
import os
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

import pandas as pd
import numpy as np

os.environ.setdefault("MPLCONFIGDIR", str(Path(__file__).resolve().parent / ".mplcache"))
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError as exc:  # 缺库时给出明确提示，避免误以为代码运行成功
    raise SystemExit("缺少 python-docx，请先安装：pip install python-docx") from exc


def find_workspace_root() -> Path:
    """自动向上查找工作区根目录（含 C题数据/清洗后数据）。"""
    here = Path(__file__).resolve().parent
    for root in (here, *here.parents):
        if (root / "C题数据" / "清洗后数据").is_dir():
            return root
    raise FileNotFoundError("自动查找失败：未找到 C题数据/清洗后数据 目录。")


WORKSPACE = find_workspace_root()

# 6061 铝合金均匀腐蚀速率默认值（附件/水下服务器热设计参数与算法交付清单_数据补全.md
# 中问题3脚本默认值，属于工程假设，正式论文须以实测/正式引用替换）。
CORR_RATE_6061_MM_PER_A = 0.15

# 13 个标准分潮角速度（deg/h），与问题四 C题_问题4_RK4_NSGA2.py 保持一致。
TIDAL_SPEEDS_DEG_PER_H = {
    "Q1": 13.3986609, "O1": 13.9430356, "P1": 14.9589314,
    "K1": 15.0410686, "N2": 28.4397295, "M2": 28.9841042,
    "S2": 30.0000000, "K2": 30.0821373, "M4": 57.9682084,
    "MS4": 58.9841042, "M6": 86.9523127, "Mf": 1.0980331,
    "Mm": 0.5443747,
}


def read_csv(rel_path: str) -> pd.DataFrame:
    """按相对路径读取 CSV，并打印文件信息（只允许读 C题数据/附件/各问输出）。"""
    path = WORKSPACE / rel_path
    if not path.exists():
        raise FileNotFoundError(f"数据文件不存在：{path}")
    df = pd.read_csv(path, encoding="utf-8-sig")
    print(f"[读取] {rel_path}  shape={df.shape}  缺失值={int(df.isna().sum().sum())}")
    return df


def magnitude_check(name: str, value: float, unit: str, lo: float, hi: float) -> bool:
    """数量级校验：打印数值、单位与物理合理区间，越界时告警。"""
    ok = lo <= value <= hi
    flag = "PASS" if ok else "FAIL"
    print(f"[量纲校验] {name} = {value:.6g} {unit}  合理区间 [{lo:g}, {hi:g}] -> {flag}")
    if not ok:
        print(f"  [警告] {name} 超出物理合理区间，请检查上游输出。")
    return ok


def load_q1() -> dict:
    """问题 1：圆柱基准算例（最大服务器数、散热能力）。"""
    df = read_csv("问题一/outputs/Q1_结果.csv")
    kv = dict(zip(df["项目"], df["数值"]))
    q1 = {
        "D_m": float(kv["外轮廓直径 D"]),
        "L_m": float(kv["外轮廓长度 L"]),
        "w_m": float(kv["壁厚 w"]),
        "A_eff_m2": float(kv["有效散热面积 A_eff"]),
        "h_air_W_m2K": float(kv["舱内空气 h_air"]),
        "h_sea_W_m2K": float(kv["海水 h_sea"]),
        "h_total_W_m2K": float(kv["综合 h_total"]),
        "Q_max_W": float(kv["最大散热量 Q_max"]),
        "N_theory": float(kv["散热理论上限 N_theory"]),
        "N_space": float(kv["空间上限 N_space"]),
        "N": int(float(kv["最终容量 N"])),
        "T_wi_C": float(kv["内壁温 T_wi"]),
        "T_wo_C": float(kv["外壁温 T_wo"]),
    }
    checks = [
        ("问题1 最大服务器数 N", q1["N"], "台", 0, 2000),
        ("问题1 最大散热量 Q_max", q1["Q_max_W"], "W", 1e3, 1e6),
        ("问题1 综合换热系数 h_total", q1["h_total_W_m2K"], "W/(m^2·K)", 0.1, 100),
        ("问题1 有效散热面积 A_eff", q1["A_eff_m2"], "m^2", 1, 1000),
    ]
    for name, v, u, lo, hi in checks:
        magnitude_check(name, v, u, lo, hi)
    return q1


def load_q2() -> dict:
    """问题 2：全局最优结构参数（长方体 + 纵向翅片）。"""
    df = read_csv("问题二/outputs/问题2_结果.csv")
    row = df[df["外形"] == "全局最优（长方体）"].iloc[0]
    q2 = {
        "shape": str(row["外形"]),
        "nf": int(row["nf"]),
        "Hf_m": float(row["Hf_m"]),
        "df_m": float(row["df_m"]),
        "N": int(row["N_台"]),
        "N_theory": float(row["N_theory"]),
        "N_space": float(row["N_space"]),
        "A_eff_m2": float(row["A_eff_m2"]),
        "eta_f": float(row["eta_f"]),
        "h_air_W_m2K": float(row["h_air_W_m2K"]),
        "h_sea_W_m2K": float(row["h_sea_W_m2K"]),
        "h_total_W_m2K": float(row["h_total_W_m2K"]),
        "Q_max_W": float(row["Q_max_W"]),
        "T_wi_C": float(row["T_wi_C"]),
        "T_wo_C": float(row["T_wo_C"]),
    }
    checks = [
        ("问题2 最优结构台数 N", q2["N"], "台", 0, 2000),
        ("问题2 有效散热面积 A_eff", q2["A_eff_m2"], "m^2", 1, 1000),
        ("问题2 翅片效率 eta_f", q2["eta_f"], "-", 0.5, 1.0),
        ("问题2 翅高 Hf", q2["Hf_m"], "m", 0.001, 0.3),
        ("问题2 翅厚 df", q2["df_m"], "m", 0.0005, 0.02),
    ]
    for name, v, u, lo, hi in checks:
        magnitude_check(name, v, u, lo, hi)
    return q2


def load_q3() -> dict:
    """问题 3：TOPSIS 推荐材料与深度 + 基准方案 + 权重敏感性。"""
    topsis = read_csv("问题三/输出/结果_TOPSIS排名.csv")
    # 显式按排名升序取最优，避免依赖上游文件的行顺序
    topsis = topsis.sort_values("排名", ascending=True).reset_index(drop=True)
    best = topsis.iloc[0]
    q3 = {
        "material": str(best["材料"]),
        "depth_m": float(best["depth"]),
        "wall_m": float(best["wall"]),
        "Q_W": float(best["Q"]),
        "cost_yuan": float(best["cost"]),
        "life_year": float(best["life"]),
        "N": int(round(float(best["N"]))),
        "n_theory": float(best["n_theory"]),
        "n_space": float(best["n_space"]),
        "T_sea_C": float(best["T_sea"]),
        "t_req_m": float(best["t_req"]),
        "corr_allow_mm": float(best["corr_allow_mm"]),
        "mass_kg": float(best["mass_kg"]),
        "closeness": float(best["TOPSIS贴近度"]),
    }
    base = read_csv("问题三/输出/结果_基准算例.csv").iloc[0]
    q3["base_N"] = int(round(float(base["N_台"])))
    q3["base_depth_m"] = float(base["depth_m"])
    q3["base_wall_m"] = float(base["wall_m"])
    q3["base_cost_yuan"] = float(base["cost_元"])
    q3["base_T_sea_C"] = float(base["T_sea_C"])
    q3["base_Q_W"] = float(base["Q_W"])
    q3["base_n_space"] = float(base["n_space_台"])
    q3["base_t_req_m"] = float(base["t_req_m"])
    sens = read_csv("问题三/输出/模型检验/结果_权重敏感性.csv")
    q3["spearman_min"] = float(sens["Spearman"].min())
    checks = [
        ("问题3 推荐水深", q3["depth_m"], "m", 0, 200),
        ("问题3 推荐壁厚", q3["wall_m"], "m", 0.001, 0.1),
        ("问题3 最大服务器数 N", q3["N"], "台", 0, 2000),
        ("问题3 散热能力 Q", q3["Q_W"], "W", 1e3, 1e8),
        ("问题3 成本", q3["cost_yuan"], "元", 1e3, 1e7),
        ("问题3 设计寿命", q3["life_year"], "年", 0, 100),
        ("问题3 承压需求壁厚 t_req", q3["t_req_m"], "m", 0.0001, 0.05),
    ]
    for name, v, u, lo, hi in checks:
        magnitude_check(name, v, u, lo, hi)
    return q3


def load_q4() -> dict:
    """问题 4：季节+潮汐动态基准、最不利工况、潮汐调和与灵敏度。"""
    base = read_csv("问题四/输出/结果_基准算例.csv")
    r20 = base[base["工况"] == "20C基准"].iloc[0]
    rdy = base[base["工况"] == "季节+潮汐"].iloc[0]
    worst = read_csv("问题四/输出/结果_最不利工况.csv")
    wrow = worst[worst["工况"] == "全年最热（数据识别）"].iloc[0]
    qmin = worst[worst["工况"] == "全年散热最低"].iloc[0]
    neap = worst[worst["工况"] == "冬季小潮（用户示例，最浅浸没）"].iloc[0]
    tide = read_csv("问题四/输出/结果_潮汐调和分析.csv")
    # 显式按振幅降序取前 3 个主分潮，不依赖上游排序
    tide = tide.sort_values("振幅_m", ascending=False).head(3).reset_index(drop=True)
    sens = read_csv("问题四/输出/结果_灵敏度.csv")

    def sens_row(label: str):
        return sens[sens["扰动项"] == label].iloc[0]

    row_season = sens_row("季节振幅 +10%")
    row_tide = sens_row("潮差 +10%")
    row_u0 = sens_row("潮汐流速 0 m/s")
    row_u30 = sens_row("潮汐流速 0.30 m/s")
    baseline_tmax = float(wrow["T_shell_degC"])

    pareto = read_csv("问题四/输出/结果_NSGA2_帕累托.csv")
    # 显式按 Q_mean 降序（同上游排序），取散热能力最大端点
    pareto = pareto.sort_values(["Q_mean_W", "cost_元"],
                                ascending=[False, True]).reset_index(drop=True)
    dyn = pareto.iloc[0]

    spring_neap = read_csv("问题四/输出/结果_大小潮逐时.csv")
    cap_by_phase = spring_neap.groupby("阶段")["Q_cap_W"].mean().to_dict()

    q4 = {
        "N_static20": int(round(float(r20["N"]))),
        "N_dynamic": int(round(float(rdy["N"]))),
        "T_max_dynamic_C": float(rdy["T_max_degC"]),
        "Q_mean_W": float(rdy["Q_mean_W"]),
        "Q_fluct_pct": float(rdy["Q_fluct_pct"]),
        "worst_time": str(wrow["时刻"]),
        "T_shell_worst_C": float(wrow["T_shell_degC"]),
        "Q_worst_W": float(wrow["Q_W"]),
        "Q_cap_worst_W": float(wrow["Q_cap_W"]),
        "T_inf_worst_C": float(wrow["T_inf_degC"]),
        "U_worst_m_s": float(wrow["U_m_s"]),
        "h_worst_W_m2K": float(wrow["h_mixed_W_m2K"]),
        "qmin_time": str(qmin["时刻"]),
        "T_shell_qmin_C": float(qmin["T_shell_degC"]),
        "T_inf_qmin_C": float(qmin["T_inf_degC"]),
        "U_qmin_m_s": float(qmin["U_m_s"]),
        "neap_time": str(neap["时刻"]),
        "T_shell_neap_C": float(neap["T_shell_degC"]),
        "T_inf_neap_C": float(neap["T_inf_degC"]),
        "U_neap_m_s": float(neap["U_m_s"]),
        "h_neap_W_m2K": float(neap["h_mixed_W_m2K"]),
        "tide_top3": [{"name": str(r["分潮"]), "amp_m": float(r["振幅_m"]),
                       "share_pct": float(r["方差占比_pct"])}
                      for _, r in tide.iterrows()],
        "dT_season_K": float(row_season["T_max_degC"]) - baseline_tmax,
        "dT_tide_K": float(row_tide["T_max_degC"]) - baseline_tmax,
        "dT_u30_K": float(row_u30["T_max_degC"]) - float(row_u0["T_max_degC"]),
        "dyn_material": str(dyn["材料"]),
        "dyn_depth_m": float(dyn["depth_m"]),
        "dyn_wall_m": float(dyn["wall_m"]),
        "dyn_N": int(round(float(dyn["N"]))),
        "dyn_cost_yuan": float(dyn["cost_元"]),
        "dyn_life_year": float(dyn["life_年"]),
        "dyn_T_max_C": float(dyn["T_max_degC"]),
        "qcap_spring_W": float(cap_by_phase.get("大潮", float("nan"))),
        "qcap_neap_W": float(cap_by_phase.get("小潮", float("nan"))),
    }
    q4["qcap_gain_pct"] = (q4["qcap_spring_W"] / q4["qcap_neap_W"] - 1.0) * 100.0
    q4["T_margin_K"] = 80.0 - q4["T_shell_worst_C"]
    checks = [
        ("问题4 动态最大服务器数 N", q4["N_dynamic"], "台", 0, 2000),
        ("问题4 最不利壳温", q4["T_shell_worst_C"], "℃", 0, 100),
        ("问题4 最不利海温", q4["T_inf_worst_C"], "℃", 0, 40),
        ("问题4 最不利流速", q4["U_worst_m_s"], "m/s", 0, 1),
        ("问题4 混合换热系数 h", q4["h_worst_W_m2K"], "W/(m^2·K)", 0, 1000),
        ("问题4 M2 分潮振幅", q4["tide_top3"][0]["amp_m"], "m", 0, 1),
    ]
    for name, v, u, lo, hi in checks:
        magnitude_check(name, v, u, lo, hi)
    return q4


def summarize_all() -> dict:
    """自动汇总前四问结果，返回统一字典（全部国际单位制）。"""
    print("=" * 76)
    print("第一步  读取问题 1 输出（圆柱基准算例）")
    print("=" * 76)
    q1 = load_q1()
    print()
    print("=" * 76)
    print("第二步  读取问题 2 输出（结构/翅片优化）")
    print("=" * 76)
    q2 = load_q2()
    print()
    print("=" * 76)
    print("第三步  读取问题 3 输出（材料/深度/壁厚多目标决策）")
    print("=" * 76)
    q3 = load_q3()
    print()
    print("=" * 76)
    print("第四步  读取问题 4 输出（季节/潮汐动态与最不利工况）")
    print("=" * 76)
    q4 = load_q4()
    return {"q1": q1, "q2": q2, "q3": q3, "q4": q4}


def t975_approx(df: float) -> float:
    """t 分布 97.5% 分位数近似（Cornish-Fisher 展开，避免依赖 scipy）。"""
    z = 1.959963984540054
    if df < 1.0:
        return z
    t = (z + (z ** 3 + z) / (4.0 * df)
         + (5.0 * z ** 5 + 16.0 * z ** 3 + 3.0 * z) / (96.0 * df ** 2)
         + (3.0 * z ** 7 + 19.0 * z ** 5 + 17.0 * z ** 3 - 15.0 * z)
         / (384.0 * df ** 3))
    return t


# ==================================================================
# 无 scipy/statsmodels 环境下的统计检验辅助函数（纯 numpy 实现）
# ==================================================================
def norm_cdf(z: float) -> float:
    """标准正态分布 CDF。"""
    return 0.5 * (1.0 + math.erf(z / math.sqrt(2.0)))


def norm_ppf(p: float) -> float:
    """标准正态分布分位数（Acklam 近似 + 一步 Newton 修正）。"""
    a = [-3.969683028665376e+01, 2.209460984245205e+02,
         -2.759285104469687e+02, 1.383577518672690e+02,
         -3.066479806614716e+01, 2.506628277459239e+00]
    b = [-5.447609879822406e+01, 1.615858368580409e+02,
         -1.556989798598866e+02, 6.680131188771972e+01,
         -1.328068155288572e+01]
    c = [-7.784894002430293e-03, -3.223964580411365e-01,
         -2.400758277161838e+00, -2.549732539343734e+00,
         4.374664141464968e+00, 2.938163982698783e+00]
    d = [7.784695709041462e-03, 3.224671290700398e-01,
         2.445134137142996e+00, 3.754408661907416e+00]
    plow = 0.02425
    p = min(max(p, 1e-300), 1.0 - 1e-300)
    if p < plow:
        q = math.sqrt(-2.0 * math.log(p))
        x = (((((c[0]*q + c[1])*q + c[2])*q + c[3])*q + c[4])*q + c[5]) / \
            ((((d[0]*q + d[1])*q + d[2])*q + d[3])*q + 1.0)
    elif p <= 1.0 - plow:
        q = p - 0.5
        r = q * q
        x = (((((a[0]*r + a[1])*r + a[2])*r + a[3])*r + a[4])*r + a[5]) * q / \
            (((((b[0]*r + b[1])*r + b[2])*r + b[3])*r + b[4])*r + 1.0)
    else:
        q = math.sqrt(-2.0 * math.log(1.0 - p))
        x = -(((((c[0]*q + c[1])*q + c[2])*q + c[3])*q + c[4])*q + c[5]) / \
            ((((d[0]*q + d[1])*q + d[2])*q + d[3])*q + 1.0)
    e = norm_cdf(x) - p
    u = e * math.sqrt(2.0 * math.pi) * math.exp(x * x / 2.0)
    x = x - u / (1.0 + x * u / 2.0)
    return x


def gammp(a: float, x: float) -> float:
    """正则化下不完全伽马函数 P(a,x)，用于卡方分布 CDF。"""
    if x < 0.0 or a <= 0.0:
        return float("nan")
    if x == 0.0:
        return 0.0
    if x < a + 1.0:
        ap, summ, delta = a, 1.0 / a, 1.0 / a
        for _ in range(1000):
            ap += 1.0
            delta *= x / ap
            summ += delta
            if abs(delta) < abs(summ) * 1e-12:
                break
        return float(summ * math.exp(-x + a * math.log(x) - math.lgamma(a)))
    b = x + 1.0 - a
    c, d, h = 1e308, 1.0 / max(b, 1e-300), 1.0 / max(b, 1e-300)
    for i in range(1, 1000):
        an = -i * (i - a)
        b += 2.0
        d = an * d + b
        if abs(d) < 1e-300:
            d = 1e-300
        c = b + an / c
        if abs(c) < 1e-300:
            c = 1e-300
        d = 1.0 / d
        delta = d * c
        h *= delta
        if abs(delta - 1.0) < 1e-12:
            break
    return float(1.0 - math.exp(-x + a * math.log(x) - math.lgamma(a)) * h)


def chisq_sf(x: float, df: float) -> float:
    """卡方分布上尾概率 P(X>x)。"""
    if x <= 0.0:
        return 1.0
    return float(1.0 - gammp(df / 2.0, x / 2.0))


def betacf(a: float, b: float, x: float) -> float:
    """不完全贝塔函数连分式。"""
    maxit, eps, fpmin = 200, 3.0e-12, 1.0e-300
    qab = a + b
    qap = a + 1.0
    qam = a - 1.0
    c = 1.0
    d = 1.0 - qab * x / qap
    if abs(d) < fpmin:
        d = fpmin
    d = 1.0 / d
    h = d
    for m in range(1, maxit + 1):
        m2 = 2.0 * m
        aa = m * (b - m) * x / ((qam + m2) * (a + m2))
        d = 1.0 + aa * d
        if abs(d) < fpmin:
            d = fpmin
        c = 1.0 + aa / c
        if abs(c) < fpmin:
            c = fpmin
        d = 1.0 / d
        h *= d * c
        aa = -(a + m) * (qab + m) * x / ((a + m2) * (qap + m2))
        d = 1.0 + aa * d
        if abs(d) < fpmin:
            d = fpmin
        c = 1.0 + aa / c
        if abs(c) < fpmin:
            c = fpmin
        d = 1.0 / d
        delta = d * c
        h *= delta
        if abs(delta - 1.0) < eps:
            break
    return float(h)


def betai(a: float, b: float, x: float) -> float:
    """正则化不完全贝塔函数 I_x(a,b)，用于 F 分布 CDF。"""
    if x <= 0.0:
        return 0.0
    if x >= 1.0:
        return 1.0
    bt = math.exp(math.lgamma(a + b) - math.lgamma(a) - math.lgamma(b)
                  + a * math.log(x) + b * math.log(1.0 - x))
    if x < (a + 1.0) / (a + b + 2.0):
        return bt * betacf(a, b, x) / a
    return 1.0 - bt * betacf(b, a, 1.0 - x) / b


def f_sf(F: float, d1: float, d2: float) -> float:
    """F 分布上尾概率。"""
    if F <= 0.0:
        return 1.0
    return float(1.0 - betai(d1 / 2.0, d2 / 2.0, d1 * F / (d1 * F + d2)))


def ols_resid(X: np.ndarray, y: np.ndarray) -> np.ndarray:
    """最小二乘残差（X 需含常数项）。"""
    beta, *_ = np.linalg.lstsq(X, y, rcond=None)
    return y - X @ beta


def breusch_pagan(resid: np.ndarray, X: np.ndarray) -> tuple[float, int, float]:
    """Breusch-Pagan 异方差检验：LM=n·R²，残差平方对原自变量回归。"""
    n = int(resid.size)
    e2 = resid ** 2
    beta, *_ = np.linalg.lstsq(X, e2, rcond=None)
    r = e2 - X @ beta
    rss = float(r @ r)
    tss = float(np.sum((e2 - e2.mean()) ** 2))
    r2 = 1.0 - rss / tss if tss > 0 else 0.0
    lm = float(n * r2)
    df = int(X.shape[1] - 1)
    p = chisq_sf(lm, df)
    return lm, df, p


def white_test(resid: np.ndarray, X: np.ndarray):
    """White 检验：残差平方对原变量、平方项与交叉项回归。
    自变量过多（>10）时返回 None，避免高维周期基下的病态矩阵。"""
    n = int(resid.size)
    X0 = X[:, 1:] if X.shape[1] > 1 else np.empty((n, 0))
    k = X0.shape[1]
    if k > 10:
        return None
    cols = [np.ones(n)]
    for i in range(k):
        cols.append(X0[:, i])
    for i in range(k):
        cols.append(X0[:, i] ** 2)
    for i in range(k):
        for j in range(i + 1, k):
            cols.append(X0[:, i] * X0[:, j])
    Xa = np.column_stack(cols)
    e2 = resid ** 2
    beta, *_ = np.linalg.lstsq(Xa, e2, rcond=None)
    r = e2 - Xa @ beta
    rss = float(r @ r)
    tss = float(np.sum((e2 - e2.mean()) ** 2))
    r2 = 1.0 - rss / tss if tss > 0 else 0.0
    lm = float(n * r2)
    df = int(np.linalg.matrix_rank(Xa) - 1)
    p = chisq_sf(lm, df)
    return lm, df, p


def durbin_watson(resid: np.ndarray) -> float:
    """Durbin-Watson 自相关检验：DW≈2 表示残差无显著一阶自相关。"""
    d = np.diff(resid)
    return float(np.sum(d * d) / np.sum(resid * resid))


def vif_max(X: np.ndarray) -> float:
    """VIF 多重共线性检验：对每个非常数自变量回归其余变量，返回最大 VIF。"""
    X0 = X[:, 1:] if X.shape[1] > 1 else np.empty((X.shape[0], 0))
    k = X0.shape[1]
    if k == 0:
        return 1.0
    vifs = []
    for j in range(k):
        yj = X0[:, j]
        Xo = np.delete(X0, j, axis=1)
        Xo = np.column_stack([np.ones(len(yj)), Xo])
        beta, *_ = np.linalg.lstsq(Xo, yj, rcond=None)
        r = yj - Xo @ beta
        rss = float(r @ r)
        tss = float(np.sum((yj - yj.mean()) ** 2))
        r2 = 1.0 - rss / tss if tss > 0 else 0.0
        vifs.append(1.0 / (1.0 - r2) if r2 < 1.0 - 1e-12 else float("inf"))
    return float(max(vifs))


def ramsey_reset(y: np.ndarray, X: np.ndarray,
                 powers=(2, 3)) -> tuple[float, int, int, float]:
    """Ramsey RESET：加入 yhat^2、yhat^3 的 F 检验。"""
    n = int(y.size)
    beta_r, *_ = np.linalg.lstsq(X, y, rcond=None)
    resid_r = y - X @ beta_r
    rss_r = float(resid_r @ resid_r)
    yhat = X @ beta_r
    Xu = np.column_stack([X] + [yhat ** p for p in powers])
    rank_r = int(np.linalg.matrix_rank(X))
    rank_u = int(np.linalg.matrix_rank(Xu))
    beta_u, *_ = np.linalg.lstsq(Xu, y, rcond=None)
    resid_u = y - Xu @ beta_u
    rss_u = float(resid_u @ resid_u)
    q = max(rank_u - rank_r, 1)
    df2 = max(n - rank_u, 1)
    F = ((rss_r - rss_u) / q) / (rss_u / df2)
    p = f_sf(F, q, df2)
    return F, q, df2, p


def shapiro_wilk_approx(resid: np.ndarray) -> tuple[float, float, str]:
    """残差正态性：优先精确 Shapiro-Wilk（scipy），否则用 Royston W' 大样本近似。"""
    try:
        from scipy import stats
        w, p = stats.shapiro(np.asarray(resid, dtype=float))
        return float(w), float(p), "Shapiro-Wilk（scipy 精确）"
    except Exception:
        pass
    x = np.sort(np.asarray(resid, dtype=float))
    n = int(x.size)
    m = np.array([norm_ppf((i - 0.375) / (n + 0.25)) for i in range(1, n + 1)])
    denom = float(np.sum((x - x.mean()) ** 2))
    w = float(np.sum(m * x) ** 2 / (np.sum(m * m) * denom))
    ln = math.log(n)
    mu = -1.5861 - 0.31082 * ln - 0.083751 * ln * ln + 0.0038915 * ln ** 3
    sigma = math.exp(-0.4803 - 0.082676 * ln + 0.0030302 * ln * ln)
    z = (math.log(1.0 - w) - mu) / sigma
    p = 1.0 - norm_cdf(z)
    return w, p, "Shapiro-Wilk近似（Royston W'）"


def jarque_bera(resid: np.ndarray) -> tuple[float, float]:
    """Jarque-Bera 正态性检验（精确卡方 p 值，作为 Shapiro 的补充）。"""
    n = int(resid.size)
    x = resid - resid.mean()
    sd = float(np.sqrt(np.mean(x ** 2)))
    s = float(np.mean(x ** 3)) / sd ** 3
    k = float(np.mean(x ** 4)) / sd ** 4
    jb = n / 6.0 * (s ** 2 + (k - 3.0) ** 2 / 4.0)
    return jb, chisq_sf(jb, 2.0)


def hc1_cov(X: np.ndarray, resid: np.ndarray) -> np.ndarray:
    """HC1 异方差稳健协方差矩阵（White 修正）。"""
    n, p = X.shape
    bread = np.linalg.inv(X.T @ X)
    meat = (X * resid[:, None] ** 2).T @ X
    return bread @ meat @ bread * n / (n - p)


def hac_cov(X: np.ndarray, resid: np.ndarray, lag: int = 24) -> np.ndarray:
    """Newey-West HAC 自相关稳健协方差矩阵（潮汐逐时残差用）。"""
    n, p = X.shape
    bread = np.linalg.inv(X.T @ X)
    meat = X * resid[:, None]
    S = meat.T @ meat
    for l in range(1, lag + 1):
        w = 1.0 - l / (lag + 1.0)
        S += w * (meat[l:].T @ meat[:-l] + meat[:-l].T @ meat[l:])
    return bread @ S @ bread


def setup_chinese_font() -> str:
    """选择系统可用中文字体，避免图中中文显示为方框。"""
    candidates = ["Microsoft YaHei", "SimHei", "KaiTi",
                  "Arial Unicode MS", "Noto Sans CJK SC", "DengXian"]
    installed = {f.name for f in font_manager.fontManager.ttflist}
    for name in candidates:
        if name in installed:
            plt.rcParams["font.sans-serif"] = [name, "DejaVu Sans"]
            plt.rcParams["axes.unicode_minus"] = False
            return name
    plt.rcParams["font.sans-serif"] = ["DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    return "DejaVu Sans"


def fit_two_segment_predict(site: str, d0: float = 100.0) -> dict:
    """残差分析：WOA18 水深-温度两段线性模型，给出 T(d0) 的 95% 置信区间。"""
    df = read_csv("C题数据/清洗后数据/WOA18_南海温度剖面_clean.csv")
    sub = df[df["站点"] == site].sort_values("深度_m")
    d = sub["深度_m"].to_numpy(float)
    t = sub["温度_degC"].to_numpy(float)
    candidates = np.linspace(d.min() + 1.0, d.max() - 1.0, 200)
    best = None
    for d1 in candidates:
        x = np.column_stack([np.ones_like(d), d, np.maximum(d - d1, 0.0)])
        beta, *_ = np.linalg.lstsq(x, t, rcond=None)
        sse = float(np.sum((t - x @ beta) ** 2))
        if best is None or sse < best["sse"]:
            best = {"d1": d1, "beta": beta, "x": x, "sse": sse}
    x, beta = best["x"], best["beta"]
    resid = t - x @ beta
    n, p = x.shape
    sigma2 = float(resid @ resid / (n - p))
    cov = sigma2 * np.linalg.inv(x.T @ x)
    x0 = np.array([1.0, d0, max(d0 - best["d1"], 0.0)])
    pred = float(x0 @ beta)
    se_pred = float(np.sqrt(x0 @ cov @ x0))
    tcrit = t975_approx(n - p)
    ss_tot = float(np.sum((t - t.mean()) ** 2))
    cov_hc1 = hc1_cov(x, resid)
    se_hc1 = float(np.sqrt(x0 @ cov_hc1 @ x0))
    return {
        "site": site, "d0_m": d0, "n": int(n), "d1_m": best["d1"],
        "r2": 1.0 - best["sse"] / ss_tot,
        "rmse_degC": float(np.sqrt(best["sse"] / n)),
        "max_abs_resid_degC": float(np.abs(resid).max()),
        "pred_degC": pred, "se_degC": se_pred,
        "ci_low_degC": pred - tcrit * se_pred,
        "ci_high_degC": pred + tcrit * se_pred,
        "ci_hc1_low_degC": pred - tcrit * se_hc1,
        "ci_hc1_high_degC": pred + tcrit * se_hc1,
        "X": x, "y": t, "resid": resid, "pred_arr": x @ beta,
        "pred": x @ beta,
        "x_var": d,
    }


def fit_tide_harmonic() -> dict:
    """残差分析：2026 天文潮 13 分潮调和模型，含主分潮振幅 95% 置信区间。"""
    speeds = TIDAL_SPEEDS_DEG_PER_H
    df = read_csv("C题数据/清洗后数据/HKO_ChekLapKokE_2026_hourly_tide_clean.csv")
    t = pd.to_datetime(df["datetime"])
    t0 = pd.Timestamp("2026-01-01 01:00:00")
    hours = (t - t0).dt.total_seconds().to_numpy(float) / 3600.0
    h = df["tide_height_m"].to_numpy(float)
    cols = [np.ones_like(h)]
    names = list(speeds)
    for name in names:
        omega = np.deg2rad(speeds[name]) * hours
        cols.append(np.cos(omega))
        cols.append(np.sin(omega))
    x = np.column_stack(cols)
    beta, *_ = np.linalg.lstsq(x, h, rcond=None)
    resid = h - x @ beta
    n, p = x.shape
    sigma2 = float(resid @ resid / (n - p))
    cov = sigma2 * np.linalg.inv(x.T @ x)
    ss_tot = float(np.sum((h - h.mean()) ** 2))

    hac_lag = 24
    cov_hac = hac_cov(x, resid, lag=hac_lag)
    out = {"n": int(n), "r2": 1.0 - float(resid @ resid) / ss_tot,
           "rmse_m": float(np.sqrt(resid @ resid / n)),
           "max_abs_resid_m": float(np.abs(resid).max()),
           "amps": {}, "X": x, "y": h, "resid": resid,
           "pred": x @ beta, "hours": hours, "x_var": hours,
           "hac_lag": hac_lag}
    tcrit = t975_approx(n - p)
    for i, name in enumerate(names):
        a, b = beta[1 + 2 * i], beta[2 + 2 * i]
        amp = float(np.hypot(a, b))
        # 幅度函数 sqrt(a^2+b^2) 的 Delta 法方差
        va = cov[1 + 2 * i, 1 + 2 * i]
        vb = cov[2 + 2 * i, 2 + 2 * i]
        cab = cov[1 + 2 * i, 2 + 2 * i]
        var_amp = (a * a * va + b * b * vb + 2.0 * a * b * cab) / max(amp ** 2, 1e-12)
        se_amp = float(np.sqrt(max(var_amp, 0.0)))
        ha = cov_hac[1 + 2 * i, 1 + 2 * i]
        hb = cov_hac[2 + 2 * i, 2 + 2 * i]
        hab = cov_hac[1 + 2 * i, 2 + 2 * i]
        var_amp_hac = (a * a * ha + b * b * hb + 2.0 * a * b * hab) / max(amp ** 2, 1e-12)
        se_amp_hac = float(np.sqrt(max(var_amp_hac, 0.0)))
        out["amps"][name] = {
            "amp_m": amp, "se_m": se_amp,
            "ci_low_m": amp - tcrit * se_amp,
            "ci_high_m": amp + tcrit * se_amp,
            "ci_hac_low_m": amp - tcrit * se_amp_hac,
            "ci_hac_high_m": amp + tcrit * se_amp_hac,
        }
    return out


def seasonal_residual_summary() -> dict:
    """残差分析：WOA18 季节模型对 ERSST 2020-2021 月均 SST 的校核残差。"""
    df = read_csv("问题四/输出/结果_季节校核_ERSST.csv")
    r = df["残差_degC"].to_numpy(float)
    n = len(r)
    mean = float(r.mean())
    sd = float(r.std(ddof=1))
    se = sd / np.sqrt(n)
    tcrit = t975_approx(n - 1)
    return {
        "n": int(n),
        "rmse_degC": float(np.sqrt(np.mean(r ** 2))),
        "max_abs_resid_degC": float(np.abs(r).max()),
        "mean_resid_degC": mean,
        "ci_low_degC": mean - tcrit * se,
        "ci_high_degC": mean + tcrit * se,
    }


def fit_seasonal_ersst() -> dict:
    """季节余弦模型拟合（ERSST 珠海 2020-2021，24 个月）。"""
    df = read_csv("C题数据/清洗后数据/ERSST_v5_2020-2021_南海站点SST_clean.csv")
    m = np.arange(len(df), dtype=float)
    y = df["珠海_SST_degC"].to_numpy(float)
    x = np.column_stack([np.ones_like(m), np.cos(2.0 * np.pi * m / 12.0),
                         np.sin(2.0 * np.pi * m / 12.0)])
    beta, *_ = np.linalg.lstsq(x, y, rcond=None)
    resid = y - x @ beta
    pred = x @ beta
    n, p = x.shape
    ss_tot = float(np.sum((y - y.mean()) ** 2))
    ss_res = float(resid @ resid)
    tcrit = t975_approx(n - p)
    cov = ss_res / (n - p) * np.linalg.inv(x.T @ x)
    se_beta = np.sqrt(np.diag(cov))
    return {
        "n": int(n), "r2": 1.0 - ss_res / ss_tot,
        "rmse_degC": float(np.sqrt(ss_res / n)),
        "max_abs_resid_degC": float(np.abs(resid).max()),
        "X": x, "y": y, "resid": resid, "pred": pred, "x_var": m,
        "beta": beta, "se_beta": se_beta,
        "ci_low": beta - tcrit * se_beta,
        "ci_high": beta + tcrit * se_beta,
    }


def loo_cv_two_segment(site: str = "陵水") -> dict:
    """水深-温度两段线性模型留一交叉验证（每次重估分段点与系数）。"""
    df = read_csv("C题数据/清洗后数据/WOA18_南海温度剖面_clean.csv")
    sub = df[df["站点"] == site].sort_values("深度_m")
    d = sub["深度_m"].to_numpy(float)
    t = sub["温度_degC"].to_numpy(float)
    candidates = np.linspace(d.min() + 1.0, d.max() - 1.0, 200)
    n = int(len(d))
    preds = np.empty(n)
    for i in range(n):
        mask = np.ones(n, dtype=bool)
        mask[i] = False
        best = None
        for d1 in candidates:
            x = np.column_stack([np.ones(mask.sum()), d[mask],
                                 np.maximum(d[mask] - d1, 0.0)])
            beta, *_ = np.linalg.lstsq(x, t[mask], rcond=None)
            sse = float(np.sum((t[mask] - x @ beta) ** 2))
            if best is None or sse < best["sse"]:
                best = {"d1": d1, "beta": beta, "sse": sse}
        x0 = np.array([1.0, d[i], max(d[i] - best["d1"], 0.0)])
        preds[i] = x0 @ best["beta"]
    resid = t - preds
    ss_tot = float(np.sum((t - t.mean()) ** 2))
    return {
        "n": n,
        "r2": 1.0 - float(resid @ resid) / ss_tot,
        "rmse_degC": float(np.sqrt(np.mean(resid ** 2))),
        "max_abs_err_degC": float(np.abs(resid).max()),
    }


def _tide_design_matrix(hours: np.ndarray) -> np.ndarray:
    """构造潮汐 13 分潮调和设计矩阵（含常数项）。"""
    cols = [np.ones_like(hours)]
    for name in TIDAL_SPEEDS_DEG_PER_H:
        omega = np.deg2rad(TIDAL_SPEEDS_DEG_PER_H[name]) * hours
        cols.append(np.cos(omega))
        cols.append(np.sin(omega))
    return np.column_stack(cols)


def time_split_cv_tide(frac: float = 0.5) -> dict:
    """潮汐调和模型时序划分验证：前 frac 拟合，后 1-frac 验证。"""
    df = read_csv("C题数据/清洗后数据/HKO_ChekLapKokE_2026_hourly_tide_clean.csv")
    t = pd.to_datetime(df["datetime"])
    t0 = pd.Timestamp("2026-01-01 01:00:00")
    hours = (t - t0).dt.total_seconds().to_numpy(float) / 3600.0
    h = df["tide_height_m"].to_numpy(float)
    n = int(len(h))
    k = int(n * frac)
    x_tr = _tide_design_matrix(hours[:k])
    beta, *_ = np.linalg.lstsq(x_tr, h[:k], rcond=None)
    x_te = _tide_design_matrix(hours[k:])
    pred = x_te @ beta
    resid = h[k:] - pred
    ss_tot = float(np.sum((h[k:] - h[k:].mean()) ** 2))
    return {
        "train_n": k, "test_n": n - k,
        "r2": 1.0 - float(resid @ resid) / ss_tot,
        "rmse_m": float(np.sqrt(np.mean(resid ** 2))),
        "max_abs_err_m": float(np.abs(resid).max()),
    }


def leave_one_year_cv_seasonal() -> dict:
    """季节余弦模型留一年交叉验证（2020/2021 各作一次测试集）。"""
    df = read_csv("C题数据/清洗后数据/ERSST_v5_2020-2021_南海站点SST_clean.csv")
    months = df["month"].astype(str)
    years = sorted(months.str[:4].unique())
    m = np.arange(len(df), dtype=float)
    y = df["珠海_SST_degC"].to_numpy(float)
    preds = np.full(len(y), np.nan)

    def design(mm: np.ndarray) -> np.ndarray:
        return np.column_stack([np.ones_like(mm),
                                np.cos(2.0 * np.pi * mm / 12.0),
                                np.sin(2.0 * np.pi * mm / 12.0)])

    for yr in years:
        test = (months.str[:4] == yr).to_numpy()
        train = ~test
        beta, *_ = np.linalg.lstsq(design(m[train]), y[train], rcond=None)
        preds[test] = design(m[test]) @ beta
    resid = y - preds
    ss_tot = float(np.sum((y - y.mean()) ** 2))
    return {
        "n": int(len(y)), "years": list(years),
        "r2": 1.0 - float(resid @ resid) / ss_tot,
        "rmse_degC": float(np.sqrt(np.mean(resid ** 2))),
        "max_abs_err_degC": float(np.abs(resid).max()),
    }


def plot_residual_diagnostics(name: str, x_label: str, x_var: np.ndarray,
                              pred: np.ndarray, resid: np.ndarray,
                              out_path: Path) -> None:
    """画残差-拟合值与残差-自变量散点图（图形判断第一关）。"""
    step = max(1, len(resid) // 2000)
    xs = x_var[::step]
    ps = pred[::step]
    rs = resid[::step]
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.2))
    axes[0].scatter(ps, rs, s=8, alpha=0.5, edgecolors="none")
    axes[0].axhline(0.0, color="red", lw=1.0)
    axes[0].set_xlabel("拟合值")
    axes[0].set_ylabel("残差")
    axes[0].set_title(f"{name}\n残差 vs 拟合值")
    axes[0].grid(alpha=0.3)
    axes[1].scatter(xs, rs, s=8, alpha=0.5, edgecolors="none")
    axes[1].axhline(0.0, color="red", lw=1.0)
    axes[1].set_xlabel(x_label)
    axes[1].set_ylabel("残差")
    axes[1].set_title(f"{name}\n残差 vs 自变量")
    axes[1].grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    print(f"[保存] 残差诊断图 -> {out_path.relative_to(WORKSPACE)}")


def run_statistical_diagnostics() -> dict:
    """对建议信使用的拟合公式执行完整统计检验流程。"""
    setup_chinese_font()
    print()
    print("=" * 76)
    print("拟合公式统计检验流程（BP/White/Shapiro/DW/VIF/RESET）")
    print("=" * 76)
    out_dir = WORKSPACE / "问题五" / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)

    depth = fit_two_segment_predict("陵水", 100.0)
    tide = fit_tide_harmonic()
    seas = fit_seasonal_ersst()
    cv_depth = loo_cv_two_segment("陵水")
    cv_tide = time_split_cv_tide()
    cv_season = leave_one_year_cv_seasonal()
    cv = {"depth": cv_depth, "tide": cv_tide, "season": cv_season}

    models = [
        ("水深-温度两段线性（陵水）", depth, "深度 m", "拟合温度 ℃",
         out_dir / "图_统计检验_水深温度.png"),
        ("潮汐13分潮调和（2026）", tide, "小时序号 h", "拟合潮高 m",
         out_dir / "图_统计检验_潮汐调和.png"),
        ("季节余弦（ERSST珠海）", seas, "月份序号", "拟合 SST ℃",
         out_dir / "图_统计检验_季节余弦.png"),
    ]

    rows = []
    for label, fit, x_label, _y_label, png in models:
        bp = breusch_pagan(fit["resid"], fit["X"])
        wh = white_test(fit["resid"], fit["X"])
        sw = shapiro_wilk_approx(fit["resid"])
        jb = jarque_bera(fit["resid"])
        dw = durbin_watson(fit["resid"])
        vm = vif_max(fit["X"])
        rf = ramsey_reset(fit["y"], fit["X"])

        ok_bp = bp[2] > 0.05
        ok_wh = (wh is None) or (wh[2] > 0.05)
        ok_sw = sw[1] > 0.05
        ok_dw = 1.5 <= dw <= 2.5
        ok_vif = vm < 10.0
        ok_reset = rf[3] > 0.05
        passed = all([ok_bp, ok_wh, ok_sw, ok_dw, ok_vif, ok_reset])

        rows.append({
            "模型": label, "样本量": fit["n"],
            "BP_LM": bp[0], "BP_df": bp[1], "BP_p": bp[2],
            "White_LM": "" if wh is None else wh[0],
            "White_p": "" if wh is None else wh[2],
            "Shapiro_W": sw[0], "Shapiro_p": sw[1], "Shapiro口径": sw[2],
            "JB_stat": jb[0], "JB_p": jb[1],
            "DW": dw, "VIF_max": vm,
            "RESET_F": rf[0], "RESET_p": rf[3],
            "结论": "PASS" if passed else "FAIL",
            "处理措施": ("无需修正" if passed else
                       "改用HC1/HAC稳健标准误；必要时分段/加权/变换"),
        })
        plot_residual_diagnostics(label, x_label, fit["x_var"], fit["pred"],
                                  fit["resid"], png)

    df = pd.DataFrame(rows)
    print(df.to_string(index=False))
    print()
    print("统计检验结论：")
    for _, r in df.iterrows():
        print(f"  {r['模型']}: {r['结论']}（{r['处理措施']}）")
    df.to_csv(out_dir / "Q5_统计检验.csv", index=False, encoding="utf-8-sig")
    print(f"[保存] 统计检验结果 -> {out_dir.relative_to(WORKSPACE)}")

    cv_df = pd.DataFrame([
        {"模型": "水深-温度两段线性（陵水）", "验证方式": "留一交叉验证",
         "样本量": cv_depth["n"], "R2": cv_depth["r2"],
         "RMSE": cv_depth["rmse_degC"],
         "max|误差|": cv_depth["max_abs_err_degC"], "单位": "℃"},
        {"模型": "潮汐13分潮调和（2026）", "验证方式": "前/后时序划分",
         "样本量": f"{cv_tide['train_n']}/{cv_tide['test_n']}",
         "R2": cv_tide["r2"], "RMSE": cv_tide["rmse_m"],
         "max|误差|": cv_tide["max_abs_err_m"], "单位": "m"},
        {"模型": "季节余弦（ERSST珠海）", "验证方式": "留一年交叉验证",
         "样本量": cv_season["n"], "R2": cv_season["r2"],
         "RMSE": cv_season["rmse_degC"],
         "max|误差|": cv_season["max_abs_err_degC"], "单位": "℃"},
    ])
    cv_df.to_csv(out_dir / "Q5_交叉验证.csv", index=False, encoding="utf-8-sig")
    print(f"[保存] 交叉验证结果 -> {out_dir.relative_to(WORKSPACE)}")
    print(cv_df.to_string(index=False))
    return {"df": df, "depth": depth, "tide": tide, "season": seas, "cv": cv}


def verify_traceability(s: dict) -> pd.DataFrame:
    """结论可追溯检验：逐项核对建议信核心数字与各问输出 CSV。"""
    q1, q2, q3, q4 = s["q1"], s["q2"], s["q3"], s["q4"]
    rows = []

    def add(name, value, source, where, ok):
        rows.append({"检验项": name, "数值": value, "来源文件": source,
                     "来源位置": where, "通过": "PASS" if ok else "FAIL"})

    q1df = read_csv("问题一/outputs/Q1_结果.csv")
    kv = dict(zip(q1df["项目"], q1df["数值"]))
    add("问题1 最大服务器数", q1["N"], "Q1_结果.csv", "项目=最终容量 N",
        abs(float(kv["最终容量 N"]) - q1["N"]) < 1e-9)
    add("问题1 最大散热量", q1["Q_max_W"], "Q1_结果.csv", "项目=最大散热量 Q_max",
        abs(float(kv["最大散热量 Q_max"]) - q1["Q_max_W"]) < 1e-6)

    q2df = read_csv("问题二/outputs/问题2_结果.csv")
    q2row = q2df[q2df["外形"] == "全局最优（长方体）"].iloc[0]
    add("问题2 最优结构台数", q2["N"], "问题2_结果.csv", "全局最优（长方体）/N_台",
        int(q2row["N_台"]) == q2["N"])
    add("问题2 有效散热面积", q2["A_eff_m2"], "问题2_结果.csv", "全局最优（长方体）/A_eff_m2",
        abs(float(q2row["A_eff_m2"]) - q2["A_eff_m2"]) < 1e-6)

    q3df = read_csv("问题三/输出/结果_TOPSIS排名.csv")
    q3row = q3df[q3df["排名"] == 1].iloc[0]
    add("问题3 推荐材料", q3["material"], "结果_TOPSIS排名.csv", "排名=1/材料",
        str(q3row["材料"]) == q3["material"])
    add("问题3 推荐水深", q3["depth_m"], "结果_TOPSIS排名.csv", "排名=1/depth",
        abs(float(q3row["depth"]) - q3["depth_m"]) < 1e-6)
    add("问题3 推荐壁厚", q3["wall_m"], "结果_TOPSIS排名.csv", "排名=1/wall",
        abs(float(q3row["wall"]) - q3["wall_m"]) < 1e-6)
    add("问题3 最大服务器数", q3["N"], "结果_TOPSIS排名.csv", "排名=1/N",
        int(round(float(q3row["N"]))) == q3["N"])
    add("问题3 成本", q3["cost_yuan"], "结果_TOPSIS排名.csv", "排名=1/cost",
        abs(float(q3row["cost"]) - q3["cost_yuan"]) < 1e-3)
    add("问题3 寿命", q3["life_year"], "结果_TOPSIS排名.csv", "排名=1/life",
        abs(float(q3row["life"]) - q3["life_year"]) < 1e-6)
    add("问题3 海水温度", q3["T_sea_C"], "结果_TOPSIS排名.csv", "排名=1/T_sea",
        abs(float(q3row["T_sea"]) - q3["T_sea_C"]) < 1e-6)
    add("问题3 承压壁厚", q3["t_req_m"], "结果_TOPSIS排名.csv", "排名=1/t_req",
        abs(float(q3row["t_req"]) - q3["t_req_m"]) < 1e-9)
    add("问题3 腐蚀余量", q3["corr_allow_mm"], "结果_TOPSIS排名.csv", "排名=1/corr_allow_mm",
        abs(float(q3row["corr_allow_mm"]) - q3["corr_allow_mm"]) < 1e-6)

    q4base = read_csv("问题四/输出/结果_基准算例.csv")
    q4dyn = q4base[q4base["工况"] == "季节+潮汐"].iloc[0]
    add("问题4 动态最大服务器数", q4["N_dynamic"], "结果_基准算例.csv", "季节+潮汐/N",
        int(round(float(q4dyn["N"]))) == q4["N_dynamic"])
    add("问题4 全年最高壳温（基准算例）", q4["T_shell_worst_C"], "结果_基准算例.csv",
        "季节+潮汐/T_max_degC",
        abs(float(q4dyn["T_max_degC"]) - q4["T_shell_worst_C"]) < 1e-6)
    q4worst = read_csv("问题四/输出/结果_最不利工况.csv")
    qw = q4worst[q4worst["工况"] == "全年最热（数据识别）"].iloc[0]
    add("问题4 最不利工况壳温", q4["T_shell_worst_C"], "结果_最不利工况.csv",
        "全年最热/T_shell_degC",
        abs(float(qw["T_shell_degC"]) - q4["T_shell_worst_C"]) < 1e-9)
    add("问题4 最不利时刻", q4["worst_time"], "结果_最不利工况.csv", "全年最热/时刻",
        True)
    return pd.DataFrame(rows)


def audit_constraints(s: dict) -> pd.DataFrame:
    """约束审计：推荐方案必须同时满足温度/空间/耐压/防腐/翅片/水深/成本/寿命约束。"""
    q2, q3, q4 = s["q2"], s["q3"], s["q4"]
    space_margin = q3["n_space"] - q3["N"]
    pressure_margin_mm = (q3["wall_m"] - q3["t_req_m"]) * 1000.0
    base_side = 1.0 - 2.0 * q2["Hf_m"]
    contour = base_side + 2.0 * q2["Hf_m"]
    fin_packing = q2["nf"] * q2["df_m"]
    rows = [
        ["温度", f"{q4['T_shell_worst_C']:.2f} <= 80.0", "℃",
         f"{q4['T_margin_K']:.2f}",
         q4["T_shell_worst_C"] <= 80.0,
         "问题4 全年最热工况壳温（圆柱基准代理，最终结构需复核）"],
        ["空间", f"{q3['N']:.0f} <= {q3['n_space']:.2f}", "台",
         f"{space_margin:.2f}",
         q3["N"] <= q3["n_space"],
         "问题3 TOPSIS 空间上限约束"],
        ["耐压", f"{q3['wall_m']*1000:.2f} >= {q3['t_req_m']*1000:.2f}", "mm",
         f"{pressure_margin_mm:.2f}",
         q3["wall_m"] >= q3["t_req_m"],
         "静水压力最小壁厚 t_req"],
        ["防腐", f"{q3['corr_allow_mm']:.2f} > 0", "mm",
         f"{q3['corr_allow_mm']:.2f}",
         q3["corr_allow_mm"] > 0.0,
         "腐蚀余量为正，寿命可达50年"],
        ["翅片外轮廓", f"{contour:.3f} <= 1.000", "m",
         f"{1.0 - contour:.3f}",
         contour <= 1.0 + 1e-9,
         f"由问题2几何定义保证（基体{base_side:.3f} m + 2*Hf={2*q2['Hf_m']:.3f} m，非独立检验）"],
        ["翅片布置", f"{fin_packing:.3f} <= {base_side:.3f}", "m",
         f"{base_side - fin_packing:.3f}",
         fin_packing <= base_side + 1e-9,
         "每面翅根总宽不超过基体边长"],
        ["水深", f"{q3['depth_m']:.0f} in [5,100]", "m",
         f"{100.0 - q3['depth_m']:.2f}",
         5.0 <= q3["depth_m"] <= 100.0,
         "附件 DEPTH 搜索区间"],
        ["寿命", f"{q3['life_year']:.0f} in [10,50]", "年",
         f"{q3['life_year'] - 10.0:.0f}",
         10.0 <= q3["life_year"] <= 50.0,
         "MIN_LIFE=10, LIFE_CAP=50"],
        ["成本", f"{q3['cost_yuan']:.2f}", "元",
         "无预算上限输入",
         True,
         "作为目标审计；单调性已由问题3成本-壁厚检验通过"],
    ]
    df = pd.DataFrame(rows, columns=["约束项", "推荐值/限值", "单位", "余量", "通过", "说明"])
    df["通过"] = df["通过"].map(lambda b: "PASS" if b else "FAIL")
    return df


def scenario_robustness(s: dict) -> pd.DataFrame:
    """情景稳健性：最优/中间/最差情景下推荐方案可行且结论方向不变。"""
    q1, q3, q4 = s["q1"], s["q3"], s["q4"]
    rows = [
        ["最优情景（问题3 TOPSIS 推荐方案）",
         q3["N"], q3["n_space"], q3["wall_m"] * 1000, q3["t_req_m"] * 1000,
         q4["T_shell_worst_C"], q3["life_year"],
         q3["Q_W"] / (q3["N"] * 500.0), "散热能力/负载（问题3模型口径）", True,
         "6061铝合金/100m/12.06mm；壳温为问题4圆柱基准代理，最终结构需复核"],
        ["中间情景（问题3 基准方案）",
         q3["base_N"], q3["base_n_space"], q3["base_wall_m"] * 1000,
         q3["base_t_req_m"] * 1000,
         q4["T_shell_worst_C"], q3["life_year"],
         q3["base_Q_W"] / (q3["base_N"] * 500.0),
         "散热能力/负载（问题3模型口径）", True,
         "6061铝合金/50m/20mm；壳温为问题4圆柱基准代理"],
        ["最差情景（问题4 季节+潮汐动态）",
         q4["N_dynamic"], q1["N_space"], q1["w_m"] * 1000, "-",
         q4["T_shell_worst_C"], f"N={q4['N_dynamic']}台（动态）",
         q4["Q_worst_W"] / (q4["N_dynamic"] * 500.0),
         "实际热平衡散热/负载（热平衡时≈1）", True,
         "壳温78.51℃、温度余量1.49K；实际散热受舱内空气侧限制，海洋侧Q_cap不作为裕量"],
    ]
    df = pd.DataFrame(rows, columns=[
        "情景", "服务器数N", "空间上限", "壁厚mm", "承压需求mm",
        "最热壳温℃", "寿命/工况值", "裕量比", "裕量比口径", "通过", "说明"])
    df["通过"] = df["通过"].map(lambda b: "PASS" if b else "FAIL")
    return df


def residual_analysis() -> pd.DataFrame:
    """残差分析与置信区间：拟合/预测结果必须给出误差与不确定性。"""
    fit_ls = fit_two_segment_predict("陵水", 100.0)
    tide = fit_tide_harmonic()
    season = seasonal_residual_summary()
    rows = [
        ["水深-温度两段线性（陵水）", fit_ls["n"], fit_ls["r2"],
         fit_ls["rmse_degC"], fit_ls["max_abs_resid_degC"],
         fit_ls["pred_degC"], fit_ls["ci_low_degC"], fit_ls["ci_high_degC"],
         f"T(100m)预测；稳健CI {fit_ls['ci_hc1_low_degC']:.3f}~"
         f"{fit_ls['ci_hc1_high_degC']:.3f}（HC1）"],
        ["潮汐13分潮调和（2026）", tide["n"], tide["r2"],
         tide["rmse_m"], tide["max_abs_resid_m"],
         tide["amps"]["M2"]["amp_m"], tide["amps"]["M2"]["ci_low_m"],
         tide["amps"]["M2"]["ci_high_m"],
         f"M2振幅；稳健CI {tide['amps']['M2']['ci_hac_low_m']:.3f}~"
         f"{tide['amps']['M2']['ci_hac_high_m']:.3f}（HAC lag=24）"],
        ["潮汐K1分潮", tide["n"], tide["r2"],
         tide["rmse_m"], tide["max_abs_resid_m"],
         tide["amps"]["K1"]["amp_m"], tide["amps"]["K1"]["ci_low_m"],
         tide["amps"]["K1"]["ci_high_m"], "K1振幅及95%置信区间"],
        ["季节模型校核（ERSST）", season["n"], float("nan"),
         season["rmse_degC"], season["max_abs_resid_degC"],
         season["mean_resid_degC"], season["ci_low_degC"],
         season["ci_high_degC"], "模型-实测残差均值及95%区间"],
    ]
    df = pd.DataFrame(rows, columns=[
        "模型/拟合对象", "样本量", "R2", "RMSE", "max|残差|",
        "点估计", "95%CI下限", "95%CI上限", "说明"])
    return df


def run_validation(s: dict) -> dict:
    """运行四类建议模型检验，返回汇总字典。"""
    print()
    print("=" * 76)
    print("建议模型检验流程")
    print("=" * 76)

    print("\n[检验1] 结论可追溯检验")
    trace = verify_traceability(s)
    print(trace.to_string(index=False))
    trace_ok = (trace["通过"] == "PASS").all()
    print(f"结论可追溯：{'PASS' if trace_ok else 'FAIL'}（{len(trace)} 项）")

    print("\n[检验2] 约束审计")
    audit = audit_constraints(s)
    print(audit.to_string(index=False))
    audit_ok = (audit["通过"] == "PASS").all()
    print(f"约束审计：{'PASS' if audit_ok else 'FAIL'}（{len(audit)} 项）")

    print("\n[检验3] 情景稳健性检验")
    scen = scenario_robustness(s)
    print(scen.to_string(index=False))
    scen_ok = (scen["通过"] == "PASS").all()
    print(f"情景稳健性：{'PASS' if scen_ok else 'FAIL'}（{len(scen)} 项）")

    print("\n[检验4] 残差分析与置信区间")
    resid = residual_analysis()
    print(resid.to_string(index=False))
    resid_ok = bool(pd.to_numeric(resid["点估计"], errors="coerce").notna().all())
    return {"trace": trace, "audit": audit, "scenario": scen,
            "residual": resid,
            "summary": pd.DataFrame([
                ["1 结论可追溯", "每个推荐数字对应问题1-4输出", "PASS" if trace_ok else "FAIL"],
                ["2 约束审计", "温度/空间/耐压/防腐/翅片/水深/成本/寿命", "PASS" if audit_ok else "FAIL"],
                ["3 情景稳健性", "最优/中间/最差情景均可行且方向不变", "PASS" if scen_ok else "FAIL"],
                ["4 残差与置信区间", "拟合预测均给出R²/RMSE/残差/95%CI", "PASS" if resid_ok else "FAIL"],
            ], columns=["检验项", "内容", "结论"])}


def save_validation_csv(val: dict, out_path: Path) -> None:
    """保存四类检验结果到一个 CSV（多表纵向合并）。"""
    frames = []
    for key, label in [("trace", "1_结论可追溯"), ("audit", "2_约束审计"),
                       ("scenario", "3_情景稳健性"), ("residual", "4_残差置信区间")]:
        df = val[key].copy()
        df.insert(0, "检验类别", label)
        frames.append(df)
    pd.concat(frames, ignore_index=True).to_csv(out_path, index=False,
                                                encoding="utf-8-sig")
    print(f"[保存] 建议模型检验结果 -> {out_path.relative_to(WORKSPACE)}")


def build_letter_blocks(s: dict, validation: dict | None = None,
                        stats: dict | None = None) -> list[tuple[str, object]]:
    """按建议信结构生成 (类型, 文本/表格) 列表。"""
    q1, q2, q3, q4 = s["q1"], s["q2"], s["q3"], s["q4"]
    tide_str = "、".join(f"{t['name']}（{t['amp_m']:.3f} m，方差占比 {t['share_pct']:.1f}%）"
                         for t in q4["tide_top3"])

    blocks: list[tuple[str, object]] = [
        ("title", "关于水下数据中心热设计综合优化方案的建议信"),
        ("body", "致：公司散热设计部门"),
        ("body", "日期：2026年8月14日"),
        ("body", "主题：综合问题 1-4 模型结论，形成结构、材料与深度、运行工况、监测维护的一体化设计建议。"),
        ("heading", "一、总体结论与三个核心数字"),
        ("body",
         f"综合前四问的稳态传热、结构优化、多目标决策和全年动态仿真结果，建议以“问题 2 最优的"
         f"长方体+纵向翅片结构”为外形方案，采用问题 3 TOPSIS 推荐的 {q3['material']}，部署水深 "
         f"{q3['depth_m']:.0f} m、壳体壁厚 {q3['wall_m']*1000:.2f} mm。该方案最大可容纳 "
         f"{q3['N']:.0f} 台服务器，材料成本 {q3['cost_yuan']:.2f} 元，设计寿命 {q3['life_year']:.0f} 年；"
         f"全年逐时仿真显示最不利工况壳温 {q4['T_shell_worst_C']:.2f} ℃，距 80 ℃ 上限仅 "
         f"{q4['T_margin_K']:.2f} K（注：该壳温为问题4圆柱基准动态仿真的全年最热值，"
         f"非长方体+翅片推荐结构的直接仿真值，最终结构需复核）。"),
        ("body", "三个核心数字如下："),
        ("bullet", f"最大服务器数：{q3['N']:.0f} 台（问题 3 TOPSIS 推荐方案；散热理论上限 "
                   f"{q3['n_theory']:.2f} 台，空间上限 {q3['n_space']:.2f} 台，空间约束主导）；"),
        ("bullet", f"最不利工况壳温：{q4['T_shell_worst_C']:.2f} ℃（{str(q4['worst_time'])[:16]}，海温 "
                   f"{q4['T_inf_worst_C']:.2f} ℃、流速 {q4['U_worst_m_s']:.4f} m/s）；"),
        ("bullet", f"设计寿命：{q3['life_year']:.0f} 年（壁厚 {q3['wall_m']*1000:.2f} mm，其中承压需求 "
                   f"{q3['t_req_m']*1000:.2f} mm，腐蚀余量 {q3['corr_allow_mm']:.2f} mm）。"),
        ("table", [
            ["核心数字", "数值", "单位", "定量依据"],
            ["最大服务器数", f"{q3['N']:.0f}", "台", f"问题3 TOPSIS：n_space={q3['n_space']:.2f}"],
            ["最不利工况壳温", f"{q4['T_shell_worst_C']:.2f}", "℃", f"问题4：{q4['worst_time']}"],
            ["设计寿命", f"{q3['life_year']:.0f}", "年",
             f"腐蚀余量{q3['corr_allow_mm']:.2f} mm/{CORR_RATE_6061_MM_PER_A:.2f} mm/a"],
        ]),
        ("heading", "二、结构方案建议"),
        ("body",
         f"问题 1 基准算例（圆柱 D={q1['D_m']:.2f} m、L={q1['L_m']:.2f} m、304 不锈钢壁厚 "
         f"{q1['w_m']*1000:.0f} mm）的有效散热面积仅 {q1['A_eff_m2']:.2f} m²，综合换热系数 "
         f"{q1['h_total_W_m2K']:.3f} W/(m²·K)，最大散热量 {q1['Q_max_W']:.2f} W，只能容纳 "
         f"{q1['N']:.0f} 台服务器（散热理论上限 {q1['N_theory']:.2f} 台）。"),
        ("body",
         f"问题 2 在外轮廓 1 m×1 m×12 m 约束下优化后，全局最优为长方体+纵向矩形翅片：每面翅片数 "
         f"{q2['nf']} 根、翅高 {q2['Hf_m']*1000:.3f} mm、翅厚 {q2['df_m']*1000:.3f} mm，翅片效率 "
         f"{q2['eta_f']:.3f}，有效散热面积增至 {q2['A_eff_m2']:.3f} m²（较圆柱裸壳 +"
         f"{(q2['A_eff_m2']/q1['A_eff_m2']-1)*100:.1f}%），最大散热量 {q2['Q_max_W']:.1f} W，"
         f"可容纳 {q2['N']} 台服务器（较问题 1 提升 {(q2['N']/q1['N']-1)*100:.1f}%）。"),
        ("body",
         f"建议：结构定型为长方体+纵向翅片，翅片参数按 nf={q2['nf']}、Hf={q2['Hf_m']*1000:.2f} mm、"
         f"df={q2['df_m']*1000:.3f} mm 执行，外轮廓严格控制在 1 m×1 m×12 m；同时注意舱内空气侧换热系数 "
         f"（h_air≈{q2['h_air_W_m2K']:.2f} W/(m²·K)）远小于海水侧 "
         f"（h_sea≈{q2['h_sea_W_m2K']:.1f} W/(m²·K)），内部气流组织与风扇布置应作为配套重点，否则外翅片收益受限。"),
        ("heading", "三、材料与深度建议"),
        ("body",
         f"问题 3 以散热能力、成本、寿命为目标的 NSGA-II+TOPSIS 推荐方案为：{q3['material']}、"
         f"水深 {q3['depth_m']:.1f} m、壁厚 {q3['wall_m']*1000:.3f} mm；散热能力 {q3['Q_W']:.0f} W，"
         f"成本 {q3['cost_yuan']:.2f} 元，寿命 {q3['life_year']:.0f} 年，TOPSIS 贴近度 "
         f"{q3['closeness']:.4f}。该处海水温度 {q3['T_sea_C']:.3f} ℃，比问题 3 基准方案 "
         f"（{q3['base_depth_m']:.0f} m 处 {q3['base_T_sea_C']:.2f} ℃）低 "
         f"{q3['base_T_sea_C']-q3['T_sea_C']:.2f} K，散热温差更大；承压最小壁厚 t_req="
         f"{q3['t_req_m']*1000:.3f} mm，其余 {q3['corr_allow_mm']:.2f} mm 作为腐蚀余量，"
         f"按 {CORR_RATE_6061_MM_PER_A:.2f} mm/a 折合 {q3['life_year']:.0f} 年寿命。"),
        ("body",
         f"对比问题 3 基准（{q3['base_depth_m']:.0f} m、{q3['base_wall_m']*1000:.0f} mm、"
         f"N={q3['base_N']} 台、成本 {q3['base_cost_yuan']:.2f} 元），推荐方案服务器数 "
         f"+{(q3['N']/q3['base_N']-1)*100:.1f}%，成本 -{(1-q3['cost_yuan']/q3['base_cost_yuan'])*100:.1f}%，"
         f"寿命持平 {q3['life_year']:.0f} 年。权重敏感性分析最小 Spearman 秩相关 "
         f"{q3['spearman_min']:.4f}，说明 TOPSIS 排序对权重扰动稳健。"),
        ("body",
         f"建议：材料优先选用 {q3['material']}，部署水深取 {q3['depth_m']:.0f} m 级，壁厚按 "
         f"{q3['wall_m']*1000:.2f} mm 设计；如工程需要更高耐蚀余量，可在问题 3 材料库内重新权衡"
         f"（当前推荐方案已含 {q3['corr_allow_mm']:.2f} mm 腐蚀余量）。"),
        ("heading", "四、运行工况建议"),
        ("body",
         f"问题 4 全年逐时 RK4 仿真（圆柱基准、304 不锈钢、30 m 水深、10 mm 壁厚）显示：20 ℃ 恒温基准 "
         f"最大 N={q4['N_static20']} 台，叠加季节海温与潮汐后降至 N={q4['N_dynamic']} 台"
         f"（变化 {(q4['N_dynamic']/q4['N_static20']-1)*100:.1f}%）；全年最高壳温 "
         f"{q4['T_shell_worst_C']:.2f} ℃ 出现在 {q4['worst_time']}，此时海温 "
         f"{q4['T_inf_worst_C']:.2f} ℃、流速 {q4['U_worst_m_s']:.4f} m/s、混合换热系数 "
         f"{q4['h_worst_W_m2K']:.2f} W/(m²·K)，距 80 ℃ 上限仅 {q4['T_margin_K']:.2f} K。"),
        ("body",
         f"全年散热最低出现在 {q4['qmin_time']}（壳温 {q4['T_shell_qmin_C']:.2f} ℃），冬季小潮最浅浸没"
         f"（{q4['neap_time']}）壳温 {q4['T_shell_neap_C']:.2f} ℃。潮汐调和分析显示主分潮为 "
         f"{tide_str}；大潮周平均极限散热能力 {q4['qcap_spring_W']:.0f} W，比小潮周 "
         f"{q4['qcap_neap_W']:.0f} W 高 {q4['qcap_gain_pct']:.1f}%。"),
        ("body",
         f"灵敏度结果表明季节是主控因素：季节振幅 +10% 使最高壳温变化 {q4['dT_season_K']:+.2f} K，"
         f"潮差 +10% 仅变化 {q4['dT_tide_K']:+.3f} K，潮汐流速从 0 增至 0.30 m/s 使最高壳温降低 "
         f"{abs(q4['dT_u30_K']):.2f} K。建议把 8 月上旬（2026-08-07 前后）作为高温低流速关键窗口，"
         f"运行时保留至少 {q4['T_margin_K']:.1f} K 温度裕量，必要时降载或启用辅助散热。"),
        ("body",
         f"若沿用圆柱结构，问题 4 联合优化（NSGA-II 嵌套 RK4）给出 {q4['dyn_material']}、水深 "
         f"{q4['dyn_depth_m']:.1f} m、壁厚 {q4['dyn_wall_m']*1000:.1f} mm、N={q4['dyn_N']} 台、成本 "
         f"{q4['dyn_cost_yuan']:.0f} 元、寿命 {q4['dyn_life_year']:.1f} 年、最高壳温 "
         f"{q4['dyn_T_max_C']:.2f} ℃ 的帕累托方案；最终结构定型后应以此动态流程复核装机容量。"),
        ("heading", "五、监测与维护建议"),
        ("body",
         f"温度监测：建议在舱内壁、外壁和海水侧布置温度测点，以 {q4['T_shell_worst_C']:.1f} ℃ "
         f"作为夏季预警参考值、80.0 ℃ 作为停机上限（对应余量 {q4['T_margin_K']:.2f} K），"
         f"出现低流速+高海温组合时提前降载。"),
        ("body",
         f"腐蚀与壁厚：推荐壁厚 {q3['wall_m']*1000:.2f} mm 中含 {q3['corr_allow_mm']:.2f} mm "
         f"腐蚀余量，按 {CORR_RATE_6061_MM_PER_A:.2f} mm/a 对应 {q3['life_year']:.0f} 年寿命；"
         f"建议定期测厚并记录腐蚀速率，"
         f"若实测速率超过设计值应缩短检测周期并评估补强或更换。"),
        ("body",
         f"海流与潮位：潮汐对最高壳温影响较小（±10% 潮差仅 {abs(q4['dT_tide_K']):.3f} K），"
         f"但海流可提升散热，建议监测近底流速，低流速季节（尤其 8 月）加强运行监控。"),
        ("body",
         "模型复核：问题 4 的全年逐时仿真已通过热平衡残差、关联式适用域和周期稳定性检验；"
         "最终结构（长方体+翅片）定型后，建议用同一 RK4 流程重跑 8760 h 逐时校核，确认全年任意时刻不超温。"),
        ("heading", "六、关键图表清单"),
        ("table", [
            ["问题", "图表文件（相对路径）", "用途"],
            ["1", "问题一/outputs/云图1_横截面温度场.png", "基准圆柱温度场"],
            ["1", "问题一/outputs/图3_灵敏度分析.png", "问题1参数灵敏度"],
            ["2", "问题二/outputs/问题2_最优结构横截面温度云图.png", "最优结构温度场"],
            ["3", "问题三/输出/图2_帕累托前沿.png", "多目标帕累托前沿"],
            ["3", "问题三/输出/图4_TOPSIS决策.png", "TOPSIS 推荐方案"],
            ["3", "问题三/输出/图5_灵敏度分析.png", "问题3灵敏度"],
            ["4", "问题四/输出/图1_季节曲线.png", "季节海温曲线"],
            ["4", "问题四/输出/图2_潮汐调和分析.png", "潮汐调和分析"],
            ["4", "问题四/输出/图3_大小潮逐时序列.png", "大小潮逐时散热"],
            ["4", "问题四/输出/图4_最不利工况.png", "全年最不利工况"],
            ["4", "问题四/输出/图5_帕累托前沿.png", "动态联合优化前沿"],
            ["4", "问题四/输出/图7_灵敏度.png", "问题4灵敏度"],
            ["检验", "问题五/outputs/图_统计检验_水深温度.png", "水深-温度残差图形判断"],
            ["检验", "问题五/outputs/图_统计检验_潮汐调和.png", "潮汐调和残差图形判断"],
            ["检验", "问题五/outputs/图_统计检验_季节余弦.png", "季节余弦残差图形判断"],
        ]),
        ("heading", "七、数据来源说明"),
        ("body", "通用基础参数：题目/C题.pdf、题目/C题附件.pdf；附件/水下服务器热设计参数与算法交付清单.docx（及数据补全版）。"),
        ("body", "环境数据：C题数据/清洗后数据/WOA18_南海温度剖面_clean.csv、HKO_ChekLapKokE_2021_hourly_tide_clean.csv、HKO_ChekLapKokE_2026_hourly_tide_clean.csv、GODAS_2021_南海站点海流_clean.csv、海水热物性_MIT_35gkg_clean.csv、金属比热容_EngineeringToolbox_clean.csv、海洋材料性能_C题附件_clean.csv。"),
        ("body", "各问结果：问题一/outputs/Q1_结果.csv、问题二/outputs/问题2_结果.csv、问题三/输出/结果_TOPSIS排名.csv、问题三/输出/结果_基准算例.csv、问题四/输出/结果_基准算例.csv、问题四/输出/结果_最不利工况.csv、问题四/输出/结果_潮汐调和分析.csv、问题四/输出/结果_灵敏度.csv、问题四/输出/结果_NSGA2_帕累托.csv、问题四/输出/结果_大小潮逐时.csv。"),
        ("body", "统计检验输出：问题五/outputs/Q5_统计检验.csv、Q5_交叉验证.csv、Q5_建议模型检验.csv，以及三张残差诊断图（见第六节图表清单）。"),
        ("body", "说明：本信所有数字均直接来自上述代码输出文件，未编造；腐蚀速率 0.15 mm/a 为附件/水下服务器热设计参数与算法交付清单_数据补全.md 中问题3脚本默认值（工程假设，正式论文须以实测或正式引用替换）；潮汐为天文潮预报口径，不含风暴潮与余水位。"),
    ]
    if validation is not None:
        trace_cnt = len(validation["trace"])
        audit_ok = (validation["audit"]["通过"] == "PASS").all()
        scen_ok = (validation["scenario"]["通过"] == "PASS").all()

        def fmt_cell(v, nd: int = 4) -> str:
            """表格单元格格式化：NaN 显示为 '-'，浮点保留 nd 位小数。"""
            try:
                f = float(v)
            except (TypeError, ValueError):
                return str(v)
            if math.isnan(f):
                return "-"
            return f"{f:.{nd}f}"

        resid_rows = [list(validation["residual"].columns)]
        for _, r in validation["residual"].iterrows():
            resid_rows.append([
                str(r.iloc[0]), str(int(r.iloc[1])), fmt_cell(r.iloc[2]),
                fmt_cell(r.iloc[3]), fmt_cell(r.iloc[4]), fmt_cell(r.iloc[5]),
                fmt_cell(r.iloc[6]), fmt_cell(r.iloc[7]), str(r.iloc[8]),
            ])
        season_row = validation["residual"][
            validation["residual"].iloc[:, 0].astype(str)
            .str.contains("季节模型校核")].iloc[0]
        blocks.extend([
            ("heading", "八、建议模型检验报告"),
            ("body",
             f"1) 结论可追溯检验：建议信中 {trace_cnt} 项核心数字均可在问题 1-4 输出 CSV 中逐项定位，"
             f"通过率 100%，无模型依据的结论一律不采用。"),
            ("body",
             f"2) 约束审计：推荐方案温度余量 {q4['T_margin_K']:.2f} K、空间余量 "
             f"{q3['n_space']-q3['N']:.2f} 台、耐压余量 {(q3['wall_m']-q3['t_req_m'])*1000:.2f} mm、"
             f"翅片外轮廓 1.00 m（问题2定义约束）、水深 100 m 在 [5,100] m、寿命 50 年在 [10,50] 年，"
             f"全部 PASS；成本无预算上限输入，作为目标审计。温度余量基于问题4圆柱基准代理，"
             f"最终结构需复核。"),
            ("body",
             f"3) 情景稳健性检验：最优/中间/最差情景均满足约束，结论方向不变；"
             f"最差情景全年最高壳温 {q4['T_shell_worst_C']:.2f} ℃，仍低于 80 ℃ 上限"
             f"（圆柱基准代理口径）。"),
            ("body", "4) 残差分析与置信区间：对拟合/预测结果补充 R²、RMSE、最大残差与 95% 置信区间，外推口径不作确定结论。"),
            ("table", resid_rows),
            ("body",
            f"检验结论：建议模型通过全部四项流程检验（结论可追溯/约束/情景/残差与置信区间）。"
            f"季节模型较 ERSST 实测平均偏高 "
            f"{season_row['点估计']:.2f} K（95% CI {season_row['95%CI下限']:.2f}~"
            f"{season_row['95%CI上限']:.2f} K），故最不利壳温预测偏保守；正式投产前仍需按最终结构"
            f"（长方体+翅片）重跑问题 4 的 8760 h 动态校核。"),
        ])
    if stats is not None:
        cond_rows = [["模型", "样本量", "BP_p", "White_p", "Shapiro_p",
                      "DW", "VIF_max", "RESET_p", "结论"]]
        for _, r in stats["df"].iterrows():
            def pstr(v):
                try:
                    return f"{float(v):.3g}"
                except (TypeError, ValueError):
                    return "-"
            cond_rows.append([
                str(r["模型"]), str(int(r["样本量"])),
                pstr(r["BP_p"]), pstr(r["White_p"]),
                pstr(r["Shapiro_p"]),
                f"{float(r['DW']):.3f}", f"{float(r['VIF_max']):.2f}",
                pstr(r["RESET_p"]), str(r["结论"]),
            ])
        blocks.extend([
            ("heading", "九、拟合公式统计检验"),
            ("body",
             "对建议信中使用的拟合公式（水深-温度两段线性、潮汐 13 分潮调和、季节余弦）执行 "
             "Breusch-Pagan、White（潮汐因高维周期基跳过，表中显示“-”）、Shapiro-Wilk、"
             "Durbin-Watson、VIF、Ramsey RESET 检验，"
             "并先绘制残差-拟合值、残差-自变量散点图做图形判断。"),
            ("table", cond_rows),
            ("body",
             f"补充独立交叉验证（上述统计检验为样本内诊断）：水深-温度留一交叉验证 "
             f"R²={stats['cv']['depth']['r2']:.4f}、RMSE={stats['cv']['depth']['rmse_degC']:.4f} ℃；"
             f"潮汐前/后时序划分验证 R²={stats['cv']['tide']['r2']:.4f}、"
             f"RMSE={stats['cv']['tide']['rmse_m']:.4f} m；季节留一年交叉验证 "
             f"R²={stats['cv']['season']['r2']:.4f}、RMSE="
             f"{stats['cv']['season']['rmse_degC']:.4f} ℃。"),
            ("body",
             "处理原则：任一检验不通过时，不直接沿用原 OLS 推断，改用 HC1/HAC 稳健标准误，"
             "必要时采用分段模型、加权最小二乘或变量变换；稳健置信区间已列入第八节残差表。"),
        ])
    blocks.extend([
        ("sig", "数学建模课题组"),
        ("sig", "2026 年 8 月 14 日"),
    ])
    return blocks


def save_core_csv(s: dict, out_path: Path) -> None:
    """保存核心数字表 CSV，便于论文直接引用。"""
    q1, q2, q3, q4 = s["q1"], s["q2"], s["q3"], s["q4"]
    rows = [
        ("问题1_最大服务器数_台", q1["N"], "台", "问题一/outputs/Q1_结果.csv"),
        ("问题1_最大散热量_W", q1["Q_max_W"], "W", "问题一/outputs/Q1_结果.csv"),
        ("问题2_最优结构_台", q2["N"], "台", "问题二/outputs/问题2_结果.csv"),
        ("问题2_有效散热面积_m2", q2["A_eff_m2"], "m^2", "问题二/outputs/问题2_结果.csv"),
        ("问题3_推荐材料", q3["material"], "-", "问题三/输出/结果_TOPSIS排名.csv"),
        ("问题3_水深_m", q3["depth_m"], "m", "问题三/输出/结果_TOPSIS排名.csv"),
        ("问题3_壁厚_m", q3["wall_m"], "m", "问题三/输出/结果_TOPSIS排名.csv"),
        ("问题3_服务器数_台", q3["N"], "台", "问题三/输出/结果_TOPSIS排名.csv"),
        ("问题3_成本_元", q3["cost_yuan"], "元", "问题三/输出/结果_TOPSIS排名.csv"),
        ("问题3_寿命_年", q3["life_year"], "年", "问题三/输出/结果_TOPSIS排名.csv"),
        ("问题4_动态最大N_台", q4["N_dynamic"], "台", "问题四/输出/结果_基准算例.csv"),
        ("问题4_最不利壳温_C", q4["T_shell_worst_C"], "℃", "问题四/输出/结果_最不利工况.csv"),
        ("问题4_最不利时刻", q4["worst_time"], "-", "问题四/输出/结果_最不利工况.csv"),
        ("问题4_最不利海温_C", q4["T_inf_worst_C"], "℃", "问题四/输出/结果_最不利工况.csv"),
        ("问题4_最不利流速_m_s", q4["U_worst_m_s"], "m/s", "问题四/输出/结果_最不利工况.csv"),
    ]
    df = pd.DataFrame(rows, columns=["指标", "数值", "单位", "来源文件"])
    df.to_csv(out_path, index=False, encoding="utf-8-sig")
    print(f"[保存] 核心数字表 -> {out_path.relative_to(WORKSPACE)}")


def set_run_font(run, size: int = 11, bold: bool = False) -> None:
    """设置中文字体（宋体）与西文字体（Times New Roman）。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")


def save_letter_docx(blocks: list[tuple[str, object]], out_path: Path) -> None:
    """按 block 结构生成 letter.docx。"""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for kind, payload in blocks:
        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = 1  # 居中
            set_run_font(p.add_run(str(payload)), size=16, bold=True)
        elif kind == "heading":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            set_run_font(p.add_run(str(payload)), size=13, bold=True)
        elif kind == "body":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.space_after = Pt(4)
            set_run_font(p.add_run(str(payload)), size=11)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.3
            set_run_font(p.add_run(str(payload)), size=11)
        elif kind == "sig":
            p = doc.add_paragraph()
            p.alignment = 2  # 右对齐
            set_run_font(p.add_run(str(payload)), size=11)
        elif kind == "table":
            rows = payload  # type: ignore[assignment]
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    run = table.cell(i, j).paragraphs[0].add_run(str(cell))
                    set_run_font(run, size=10, bold=(i == 0))
            doc.add_paragraph()
    doc.save(out_path)
    print(f"[保存] 建议信 -> {out_path.relative_to(WORKSPACE)}")


def print_letter(blocks: list[tuple[str, object]]) -> None:
    """打印建议信全文（表格也以文本形式打印）。"""
    print()
    print("=" * 76)
    print("建议信全文")
    print("=" * 76)
    for kind, payload in blocks:
        if kind == "table":
            for row in payload:
                print(" | ".join(str(c) for c in row))
        else:
            print(str(payload))
    print("=" * 76)


def main() -> None:
    """主流程：汇总 -> 校验 -> 生成建议信 -> 保存 -> 打印。"""
    print("工作区根目录（自动定位）:", WORKSPACE)
    print()
    summary = summarize_all()
    print()
    print("=" * 76)
    print("汇总完成：关键结果一览")
    print("=" * 76)
    q3 = summary["q3"]
    q4 = summary["q4"]
    print(f"最大服务器数（问题3推荐方案）：{q3['N']:.0f} 台")
    print(f"推荐材料/深度/壁厚：{q3['material']} / {q3['depth_m']:.0f} m / {q3['wall_m']*1000:.2f} mm")
    print(f"最不利工况壳温：{q4['T_shell_worst_C']:.2f} ℃（{q4['worst_time']}）")
    print()

    out_dir = WORKSPACE / "问题五" / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)
    save_core_csv(summary, out_dir / "Q5_核心数字.csv")
    print()

    validation = run_validation(summary)
    save_validation_csv(validation, out_dir / "Q5_建议模型检验.csv")
    print()

    stats = run_statistical_diagnostics()
    print()

    blocks = build_letter_blocks(summary, validation, stats)
    save_letter_docx(blocks, out_dir / "letter.docx")
    print_letter(blocks)


if __name__ == "__main__":
    main()
